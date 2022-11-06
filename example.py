# ПРИМЕРЫ ИСПОЛЬЗОВАНИЯ ПАКЕТОВ PYTHON

#============================================================================================

# pprint для вывода больших списков и словарей в столбик
from pprint import pprint
pprint(dir(__builtins__))

#============================================================================================

# random для генерации случайных значений
import random

my_favorite_songs = [
    'Waste a Moment',
    'New Salvation', 
    'Staying\' Alive', 
    'Out of Touch', 
    'A Sorta Fairytale', 
    'Easy', 
    'Beautiful Day', 
    'Nowhere to Run', 
    'In This World'
    ]

print("Выбор случайной песни из списка - ", random.choice(my_favorite_songs))
print("Три случайных уникальных песни - ", random.sample(my_favorite_songs, 3))
print("Три случайных песни - ", random.choices(my_favorite_songs, k=3))
print("Cлучайное число от 0 до 5: ", random.randint(0, 5))

print("Генерация случайного числа в пределах заданного промежутка")
print(random.randrange(100, 10000))

#============================================================================================

# time для управления временем
# вывод каждый 5 секунд
import time

def hello():
    print('Привет')


# while True:
#     hello()
#     time.sleep(5) 

#============================================================================================

# os для управления файлами и каталогами
import os
#### Получить абсолютные пути только файлов. 
def get_paths(path = '.'):

    for name in os.listdir(path):
        abs_path = os.path.abspath(os.path.join(path, name))
        if os.path.isfile(abs_path) is True:
            yield abs_path
        
        elif os.path.isdir(abs_path) is True:
            yield from get_paths(abs_path)


for i in get_paths('Хранилище'):
    print(i)

#============================================================================================

# pandas для обработки табличных данных
# conda install pandas
# conda install openpyxl
import pandas as pd
import os

def char_clean(s):
    s = str(s) 
    new_s = '' 
    for i in s:
        if i.isdigit() is True:
            new_s += i 
        else:
            pass  
    return new_s 

PATH = f'{os.getcwd()}/phone_numbers.xlsx'
NEWFILE_PATH = f'{os.getcwd()}/clean_phone_numbers.xlsx'

excel_df = pd.read_excel(PATH, sheet_name='Worksheet')

phone_list = excel_df['phone_number'].values.tolist()
clean_phone_list = [] 
phone = ''

for i in range(len(phone_list)):
    phone = phone_list[i] 
    new_phone = char_clean(phone)  
    clean_phone_list.append(new_phone) 

print('\n', '*'*77, '\nВот ваш результат:\n')
print(clean_phone_list)


new_df = pd.DataFrame(clean_phone_list, columns =['phones'])
new_df.to_excel(NEWFILE_PATH, sheet_name='clean_numbers', index=False)

#============================================================================================

# docx для обработки текстовых файлов .docx
# conda install docx
import docx
import os

def folder_check(path, folder):
    if folder in os.listdir(path):
        pass
    else:
        os.mkdir(f'{path}/{folder}')

def change_font(dirs):
    for dir in dirs:
        doc = docx.Document(dir)

        filename, file_extension = os.path.splitext(os.path.basename(dir))

        for paragraph in doc.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = 1.5

            for run in paragraph.runs:
                font = run.font
                run.font.size = docx.shared.Pt(14)
                run.font.name = 'Times New Roman'

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.line_spacing = 1.5

                        for run in paragraph.runs:
                            font = run.font
                            run.font.size = docx.shared.Pt(14)
                            run.font.name = 'Times New Roman'

        doc.save(f'{os.getcwd()}/Итог/{filename}{file_extension}')

paths = []
folder = os.getcwd()
os.chdir('Исходники')

folder_check(f'{folder}/Исходники', 'Итог')

for root, dirs, files in os.walk(folder):
    for file in files:
        if file.endswith('docx') and not file.startswith('~'):
            paths.append(os.path.join(root, file))

change_font(paths)

#============================================================================================

# Фреймворк flask ля создания веб-сайтов
# conda install flask
from flask import Flask
app = Flask(__name__)

@app.route("/")
def hello():
    return "Hello, World!"
